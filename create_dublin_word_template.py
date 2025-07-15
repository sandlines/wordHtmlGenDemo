#!/usr/bin/env python3
"""
Script to create Dublin Word template that matches the HTML structure
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

def create_dublin_word_template():
    """Create Dublin Word template matching the HTML structure"""
    
    # Create new document
    doc = Document()
    
    # Set document margins (0.75 inches like HTML)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # HEADER SECTION - Three columns
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Remove table borders
    for row in header_table.rows:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'))
    
    # Left column - Councilmembers
    left_cell = header_table.cell(0, 0)
    left_cell.width = Inches(2.5)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add councilmembers title
    title_run = left_para.add_run("COUNCILMEMBERS")
    title_run.font.size = Pt(11)
    title_run.font.name = "Times New Roman"
    title_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray color
    title_run.bold = True
    
    # Add councilmembers list
    councilmembers = [
        "Dr. Sherry Hu, Mayor",
        "Kashef Qaadri, Vice Mayor", 
        "Jean Josey, Councilmember",
        "Michael McCorriston, Councilmember",
        "John Morada, Councilmember"
    ]
    
    for member in councilmembers:
        member_para = left_cell.add_paragraph()
        member_run = member_para.add_run(member)
        member_run.font.size = Pt(10)
        member_run.font.name = "Times New Roman"
    
    # Center column - Dublin logo and branding
    center_cell = header_table.cell(0, 1)
    center_cell.width = Inches(2.5)
    center_para = center_cell.paragraphs[0]
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add logo placeholder (shamrock symbol) - make it more prominent
    logo_para = center_cell.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_para.add_run("☘")
    logo_run.font.size = Pt(48)  # Larger shamrock
    logo_run.font.color.rgb = RGBColor(46, 139, 87)  # Forest green color
    
    # Add city name
    city_para = center_cell.add_paragraph()
    city_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city_run = city_para.add_run("DUBLIN")
    city_run.font.size = Pt(24)
    city_run.font.name = "Times New Roman"
    city_run.font.color.rgb = RGBColor(46, 139, 87)  # Forest green color
    city_run.bold = True
    
    # Add state name
    state_para = center_cell.add_paragraph()
    state_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    state_run = state_para.add_run("CALIFORNIA")
    state_run.font.size = Pt(10)
    state_run.font.name = "Times New Roman"
    state_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray color
    
    # Right column - Location info
    right_cell = header_table.cell(0, 2)
    right_cell.width = Inches(2.5)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add location information
    location_info = [
        "Peter W. Snyder Council Chamber",
        "Dublin Civic Center",
        "100 Civic Plaza",
        "Dublin, CA 94568",
        "www.dublin.ca.gov"
    ]
    
    for info in location_info:
        info_para = right_cell.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        info_run = info_para.add_run(info)
        info_run.font.size = Pt(10)
        info_run.font.name = "Times New Roman"
    
    # Add spacing after header
    doc.add_paragraph()
    
    # MEETING TITLE SECTION
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Regular Meeting of the")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = "Times New Roman"
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray color
    
    # Main title
    main_title_para = doc.add_paragraph()
    main_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_title_run = main_title_para.add_run("DUBLIN CITY COUNCIL")
    main_title_run.font.size = Pt(28)
    main_title_run.font.name = "Times New Roman"
    main_title_run.bold = True
    
    # Add spacing
    doc.add_paragraph()
    
    # MEETING DETAILS - Date and Location
    details_table = doc.add_table(rows=1, cols=2)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Remove table borders
    for row in details_table.rows:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'))
    
    # Date (left)
    date_cell = details_table.cell(0, 0)
    date_para = date_cell.paragraphs[0]
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_run = date_para.add_run("{{meeting_date}}")
    date_run.font.size = Pt(12)
    date_run.font.name = "Times New Roman"
    date_run.bold = True
    
    # Location (right)
    location_cell = details_table.cell(0, 1)
    location_para = location_cell.paragraphs[0]
    location_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    location_run = location_para.add_run("Location: Peter W. Snyder")
    location_run.font.size = Pt(12)
    location_run.font.name = "Times New Roman"
    location_run.bold = True
    
    # Add more location lines
    for line in ["Council Chamber", "100 Civic Plaza", "Dublin, CA 94568"]:
        loc_para = location_cell.add_paragraph()
        loc_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        loc_run = loc_para.add_run(f"                    {line}")
        loc_run.font.size = Pt(12)
        loc_run.font.name = "Times New Roman"
    
    # MEETING TIME
    time_para = doc.add_paragraph()
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_run = time_para.add_run("REGULAR MEETING 7:00 PM")
    time_run.font.size = Pt(14)
    time_run.font.name = "Times New Roman"
    time_run.bold = True
    
    # Add spacing
    doc.add_paragraph()
    
    # PROCEDURES BOX - Create a bordered table to match HTML version
    procedures_table = doc.add_table(rows=1, cols=1)
    procedures_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    procedures_table.autofit = False
    
    # Set table width to full page width
    procedures_table.columns[0].width = Inches(6.5)
    
    # Add border to the table
    for row in procedures_table.rows:
        for cell in row.cells:
            # Add thick border around the cell
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="thick" w:sz="24" w:space="0" w:color="000000"/><w:left w:val="thick" w:sz="24" w:space="0" w:color="000000"/><w:bottom w:val="thick" w:sz="24" w:space="0" w:color="000000"/><w:right w:val="thick" w:sz="24" w:space="0" w:color="000000"/></w:tcBorders>'))
    
    # Add content to the procedures box
    procedures_cell = procedures_table.cell(0, 0)
    
    # Add title
    procedures_title_para = procedures_cell.paragraphs[0]
    procedures_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    procedures_title_run = procedures_title_para.add_run("Additional Meeting Procedures")
    procedures_title_run.font.size = Pt(14)
    procedures_title_run.font.name = "Times New Roman"
    procedures_title_run.bold = True
    procedures_title_run.underline = True
    
    # Add procedures content inside the bordered table
    proc_content = [
        "This City Council meeting will be broadcast live on Comcast T.V. channel 28 beginning at 7:00 p.m. This meeting will also be livestreamed at www.tv30.org and on the City's website at: https://dublin.ca.gov/ccmeetings",
        "",
        "For the convenience of the City and as a courtesy to the public, members of the public who wish to offer comments electronically have the option of giving public comment via Zoom, subject to the following procedures:",
        "",
        "• Fill out an online speaker slip available at www.dublin.ca.gov. The speaker slip will be made available at 10:00 a.m. on Tuesday, May 6, 2025. Upon submission, you will receive Zoom link information from the City Clerk. Speakers slips will be accepted until the staff presentation ends, or until the public comment period on non-agenda items is closed.",
        "",
        "• Once connected to the Zoom platform using the Zoom link information from the City Clerk, the public speaker will be added to the Zoom webinar as an attendee and muted. The speaker will be able to observe the meeting from the Zoom platform.",
        "",
        "• When the agenda item upon which the individual would like to comment is addressed, the City Clerk will announce the speaker in the meeting when it is their time to give public comment. The speaker will then be unmuted to give public comment via Zoom.",
        "",
        "• Technical difficulties may occur that make the option unavailable, and, in such event, the meeting will continue despite the inability to provide the option."
    ]
    
    for content in proc_content:
        if content:
            proc_para = procedures_cell.add_paragraph()
            proc_run = proc_para.add_run(content)
            proc_run.font.size = Pt(10)
            proc_run.font.name = "Times New Roman"
        else:
            procedures_cell.add_paragraph()
    
    # Add spacing
    doc.add_paragraph()
    
    # AGENDA ITEMS SECTION
    agenda_title_para = doc.add_paragraph()
    agenda_title_run = agenda_title_para.add_run("AGENDA ITEMS")
    agenda_title_run.font.size = Pt(16)
    agenda_title_run.font.name = "Times New Roman"
    agenda_title_run.bold = True
    
    # Add placeholder for agenda items
    agenda_content_para = doc.add_paragraph()
    agenda_content_run = agenda_content_para.add_run("{{agenda_content}}")
    agenda_content_run.font.size = Pt(12)
    agenda_content_run.font.name = "Times New Roman"
    
    # Save the template
    template_path = "templates/dublin-agenda-word-template.docx"
    doc.save(template_path)
    
    print(f"✅ Dublin Word template created: {template_path}")
    return template_path

if __name__ == "__main__":
    create_dublin_word_template() 