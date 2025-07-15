from flask import Flask, request, jsonify, send_file, render_template_string
from flask_cors import CORS
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Pt
from docx import Document
import os
import subprocess
import json
import tempfile
import uuid
from datetime import datetime
import threading
import time
from jinja2 import Template

app = Flask(__name__)
CORS(app)

# Store active sessions and their state
sessions = {}

def format_agenda_sections(sections_data):
    """Format agenda sections from form data"""
    formatted_sections = []
    for section in sections_data:
        if section.get('type') == 'break':
            # Handle section breaks with enhanced formatting
            break_title = section.get('title', '────────────────────────────────────────')
            # Add centered section break with visual separators
            formatted_sections.append("") # Empty line before
            formatted_sections.append("═" * 60) # Top border
            formatted_sections.append(f"{'═' * 10} {break_title.center(36)} {'═' * 10}")
            formatted_sections.append("═" * 60) # Bottom border
            formatted_sections.append("") # Empty line after
        else:
            # Handle regular sections
            section_number = section.get('number', '')
            section_title = section.get('title', '')
            if section_number:
                section_text = f"{section_number}. {section_title}"
            else:
                section_text = section_title
            formatted_sections.append(section_text)
            formatted_sections.append("")
            
            for item in section.get('items', []):
                if item.get('prefix') or item.get('number'):
                    item_text = f"     {item.get('prefix', '')}{item.get('number', '')} {item['title']}"
                else:
                    item_text = f"     {item['title']}"
                formatted_sections.append(item_text)
                
                if item.get('attachments'):
                    for attachment in item['attachments']:
                        formatted_sections.append(f"          Attachment: {attachment}")
            
            formatted_sections.append("")
    
    return "\n".join(formatted_sections)

def apply_section_break_formatting(doc, agenda_content):
    """Apply enhanced formatting to section breaks in the document"""
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Check if this is a section break line
            if text.startswith('═') and len(text) >= 60:
                # Center the section break
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Apply formatting to the section break
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    
                    # If this is the title line (middle line with text), make it larger
                    if len(text) > 60 and '═' in text and not text.replace('═', '').replace(' ', '').isdigit():
                        run.font.size = Pt(14)
                        run.font.bold = True
        
        return True
        
    except Exception as e:
        print(f"Section break formatting failed: {e}")
        return False

def apply_font_customization(docx_filename, font_settings):
    """Apply font customization and margins to the generated Word document"""
    try:
        # Open the document
        doc = Document(docx_filename)
        
        # Get font settings with defaults
        document_font = font_settings.get('document_font', 'Times New Roman')
        heading_font = font_settings.get('heading_font', 'Times New Roman')
        font_size = font_settings.get('font_size', 12)
        heading_size = font_settings.get('heading_size', 18)
        
        # Get margin settings with defaults (in inches)
        margin_top = font_settings.get('margin_top', 1)
        margin_bottom = font_settings.get('margin_bottom', 1)
        margin_left = font_settings.get('margin_left', 1)
        margin_right = font_settings.get('margin_right', 1)
        
        # Apply margins to all sections
        from docx.shared import Inches
        for section in doc.sections:
            section.top_margin = Inches(margin_top)
            section.bottom_margin = Inches(margin_bottom)
            section.left_margin = Inches(margin_left)
            section.right_margin = Inches(margin_right)
        
        # Apply fonts to all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only modify non-empty paragraphs
                for run in paragraph.runs:
                    if run.font.size and run.font.size >= Pt(16):
                        # This is likely a heading
                        run.font.name = heading_font
                        run.font.size = Pt(heading_size)
                    else:
                        # This is regular text
                        run.font.name = document_font
                        run.font.size = Pt(font_size)
        
        # Apply fonts to tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.size and run.font.size >= Pt(16):
                                # This is likely a heading
                                run.font.name = heading_font
                                run.font.size = Pt(heading_size)
                            else:
                                # This is regular text
                                run.font.name = document_font
                                run.font.size = Pt(font_size)
        
        # Save the modified document
        doc.save(docx_filename)
        
        return True
        
    except Exception as e:
        print(f"Font customization failed: {e}")
        return False

def generate_word_document(template_data):
    """Generate Word document from template data"""
    template_path = "templates/comprehensive_agenda_template.docx"
    logo_path = "assets/sausalito.jpeg"
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    if not os.path.exists(logo_path):
        raise FileNotFoundError(f"Logo not found: {logo_path}")
    
    tpl = DocxTemplate(template_path)
    city_logo = InlineImage(tpl, logo_path, width=Mm(30), height=Mm(30))
    
    # Process the template data
    context = {
        "city_name": template_data.get("city_name", "City Name"),
        "city_logo": city_logo,
        "meeting_type": template_data.get("meeting_type", "Regular Meeting"),
        "location": {"address": template_data.get("address", "City Hall")},
        "meeting_date": template_data.get("meeting_date", "Date TBD"),
        "special_time": template_data.get("special_time", "6:00 PM"),
        "regular_time": template_data.get("regular_time", "7:00 PM"),
        "zoom": {
            "url": template_data.get("zoom_url", "https://zoom.us/j/123456789"),
            "passcode": template_data.get("zoom_passcode", "123456"),
            "phone_list": template_data.get("zoom_phone", "+1 669 900 6833")
        },
        "agenda_content": format_agenda_sections(template_data.get("agenda_sections", [])),
        "lead_department": {
            "name": template_data.get("department_name", "Administration Department"),
            "address": template_data.get("address", "City Hall"),
            "phone": template_data.get("department_phone", "(555) 123-4567"),
            "email": template_data.get("department_email", "admin@city.gov")
        },
        "council_list": template_data.get("council_members", ""),
        "staff_list": template_data.get("staff_list", ""),
    }
    
    tpl.render(context)
    
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tpl.save(temp_file.name)
    
    # Apply font customization if provided
    font_settings = template_data.get('font_settings', {})
    if font_settings:
        apply_font_customization(temp_file.name, font_settings)
    
    # Apply section break formatting
    try:
        doc = Document(temp_file.name)
        apply_section_break_formatting(doc, context["agenda_content"])
        doc.save(temp_file.name)
    except Exception as e:
        print(f"Section break formatting failed: {e}")
    
    return temp_file.name

def generate_dublin_word_document(template_data):
    """Generate Dublin Word document from template data"""
    template_path = "templates/dublin-agenda-word-template.docx"
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Dublin Word template not found: {template_path}")
    
    try:
        from docxtpl import DocxTemplate
        
        # Load the template
        tpl = DocxTemplate(template_path)
        
        # Process the template data
        context = {
            "meeting_date": template_data.get("meeting_date", "Date TBD"),
            "agenda_content": format_agenda_sections(template_data.get("agenda_sections", []))
        }
        
        # Render the template
        tpl.render(context)
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.close()
        
        # Save the document
        tpl.save(temp_file.name)
        
        # Apply font customization if provided
        font_settings = template_data.get('font_settings', {})
        if font_settings:
            apply_font_customization(temp_file.name, font_settings)
        
        return temp_file.name
        
    except Exception as e:
        raise Exception(f"Dublin Word document generation failed: {str(e)}")

def convert_to_pdf(docx_filename):
    """Convert Word to PDF using LibreOffice"""
    try:
        output_dir = os.path.dirname(docx_filename)
        result = subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf',
            '--outdir', output_dir, docx_filename
        ], capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            pdf_filename = docx_filename.replace('.docx', '.pdf')
            if os.path.exists(pdf_filename):
                return pdf_filename
        
        raise Exception(f"LibreOffice conversion failed: {result.stderr}")
        
    except Exception as e:
        raise Exception(f"PDF conversion failed: {str(e)}")

@app.route('/')
def index():
    """Serve the main template editor interface"""
    try:
        with open('demo_frontend.html', 'r') as f:
            return f.read()
    except FileNotFoundError:
        return "Template Editor Demo - Frontend file not found"

def generate_dublin_html_document(template_data):
    """Generate HTML document for Dublin agenda using HTML template"""
    template_path = "templates/dublin-agenda-template.html"
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    # Read the HTML template
    with open(template_path, 'r') as f:
        template_content = f.read()
    
    # Create Jinja2 template
    template = Template(template_content)
    
    # Prepare agenda items
    agenda_items = []
    for section in template_data.get("agenda_sections", []):
        if section.get('type') == 'section':
            for item in section.get('items', []):
                agenda_items.append({
                    'title': item.get('title', ''),
                    'description': item.get('description', ''),
                    'time_estimate': item.get('time_estimate', '')
                })
    
    # Render the template
    html_content = template.render(
        meeting_date=template_data.get("meeting_date", "Date TBD"),
        agenda_items=agenda_items
    )
    
    return html_content

def generate_dublin_tiptap_document(template_data):
    """Generate Dublin agenda cover using TipTap content"""
    try:
        # Get TipTap content from the request
        tiptap_content = template_data.get('tiptap_content', '')
        
        # Additional meeting data for dynamic content
        meeting_data = {
            'date': template_data.get('meeting_date', 'Date TBD'),
            'city': template_data.get('city_name', 'Dublin'),
            'time': template_data.get('meeting_time', '7:00 PM')
        }
        
        # For now, generate a simple HTML document with TipTap content
        # In a real implementation, you would use the TypeScript utility
        html_content = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Dublin City Council Agenda</title>
            <style>
                {get_dublin_styles()}
            </style>
        </head>
        <body>
            <div class="dublin-cover-page">
                {convert_tiptap_to_html(tiptap_content)}
            </div>
        </body>
        </html>
        """
        
        return html_content
        
    except Exception as e:
        raise Exception(f"Dublin TipTap document generation failed: {str(e)}")

def get_dublin_styles():
    """Get Dublin-specific CSS styles"""
    return """
        @page {
            size: letter;
            margin: 0.75in;
        }
        
        body {
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.3;
            color: #000;
            margin: 0;
            padding: 0;
            background: white;
        }

        .dublin-cover-page {
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.3;
            color: #000;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 0;
            background: white;
        }

        /* Attribute-based paragraph styling */
        p[data-align="left"] {
            text-align: left;
        }

        p[data-align="center"] {
            text-align: center;
        }

        p[data-align="right"] {
            text-align: right;
        }

        p[data-spacing="tight"] {
            line-height: 1.15;
            margin-bottom: 4pt;
        }

        p[data-spacing="normal"] {
            line-height: 1.3;
            margin-bottom: 12pt;
        }

        p[data-spacing="loose"] {
            line-height: 1.5;
            margin-bottom: 20pt;
        }

        p[data-variant="fine-print"] {
            font-size: 9pt;
            color: #666;
        }

        p[data-variant="heading"] {
            font-size: 14pt;
            font-weight: bold;
        }

        p[data-variant="subtitle"] {
            font-size: 11pt;
            color: #666;
        }

        .council-list {
            text-align: left;
        }

        .council-title {
            font-size: 11pt;
            color: #666;
            font-weight: 500;
            margin-bottom: 8px;
            text-transform: uppercase;
            font-family: 'Times New Roman', Times, serif;
        }

        .council-member {
            font-size: 10pt;
            color: #333;
            margin-bottom: 2px;
            line-height: 1.3;
            font-family: 'Times New Roman', Times, serif;
        }

        .council-members-content {
            margin-top: 8px;
        }

        .council-members-content p {
            font-size: 10pt;
            color: #333;
            margin-bottom: 2px;
            line-height: 1.3;
            font-family: 'Times New Roman', Times, serif;
        }

        .dublin-logo-container {
            text-align: center;
            margin: 1rem 0;
        }

        .dublin-logo {
            display: block;
            margin: 0 auto;
            width: 80px;
            height: 80px;
        }

        .location-block {
            text-align: right;
            font-size: 10pt;
            color: #666;
            line-height: 1.2;
        }

        .location-line {
            margin-bottom: 2px;
            font-family: 'Times New Roman', Times, serif;
        }

        .location-block p {
            font-size: 10pt;
            color: #666;
            margin-bottom: 2px;
            line-height: 1.2;
            font-family: 'Times New Roman', Times, serif;
            text-align: right;
        }

        .notice-box {
            border: 2px solid #000;
            padding: 20px;
            margin: 25px 0;
            page-break-inside: avoid;
        }

        .notice-box-title {
            text-align: center;
            font-size: 14pt;
            font-weight: bold;
            text-decoration: underline;
            margin-bottom: 15px;
            font-family: 'Times New Roman', Times, serif;
        }

        .notice-box-content {
            font-size: 10pt;
            line-height: 1.4;
            text-align: left;
        }

        .section-break {
            text-align: center;
            font-size: 14pt;
            font-weight: bold;
            margin: 25px 0;
            font-family: 'Times New Roman', Times, serif;
        }

        .dublin-title {
            font-size: 28pt;
            font-weight: bold;
            color: #2e8b57;
            text-align: center;
            margin: 1rem 0;
            font-family: 'Times New Roman', Times, serif;
        }

        .dublin-title-main {
            font-size: 28pt;
            font-weight: bold;
            color: #2e8b57;
            text-align: center;
            margin: 1rem 0;
            font-family: 'Times New Roman', Times, serif;
        }

        .dublin-title-sub {
            font-size: 14pt;
            color: #666;
            text-align: center;
            margin-bottom: 25px;
            font-family: 'Times New Roman', Times, serif;
        }
    """

def convert_tiptap_to_html(tiptap_content):
    """Convert TipTap JSON content to HTML"""
    if not tiptap_content:
        return '<p>No content provided</p>'
    
    # For now, return the content as-is if it's already a string
    if isinstance(tiptap_content, str):
        return tiptap_content
    
    # If it's JSON, try to convert it to a simple HTML representation
    try:
        if isinstance(tiptap_content, dict) and tiptap_content.get('type') == 'doc':
            html_parts = []
            for node in tiptap_content.get('content', []):
                html_parts.append(convert_tiptap_node_to_html(node))
            return ''.join(html_parts)
        else:
            return str(tiptap_content)
    except Exception:
        return str(tiptap_content)

def convert_tiptap_node_to_html(node):
    """Convert a single TipTap node to HTML"""
    if not node:
        return ''
    
    node_type = node.get('type', '')
    attrs = node.get('attrs', {})
    content = node.get('content', [])
    
    if node_type == 'dublinParagraph':
        attr_str = ''
        if attrs.get('align') != 'left':
            attr_str += f' data-align="{attrs.get("align")}"'
        if attrs.get('spacing') != 'normal':
            attr_str += f' data-spacing="{attrs.get("spacing")}"'
        if attrs.get('variant') != 'body':
            attr_str += f' data-variant="{attrs.get("variant")}"'
        if attrs.get('color'):
            attr_str += f' data-color="{attrs.get("color")}"'
        
        text_content = ''.join(convert_tiptap_node_to_html(child) for child in content)
        return f'<p{attr_str}>{text_content}</p>'
    
    elif node_type == 'locationBlock':
        content_html = ''.join(convert_tiptap_node_to_html(child) for child in content)
        return f'<div class="location-block">{content_html}</div>'
    
    elif node_type == 'councilList':
        content_html = ''.join(convert_tiptap_node_to_html(child) for child in content)
        return f'<div class="council-list"><h3 class="council-title">COUNCILMEMBERS</h3><div class="council-members-content">{content_html}</div></div>'

    elif node_type == 'coverHeader':
        content_html = ''.join(convert_tiptap_node_to_html(child) for child in content)
        return f'<div class="cover-header">{content_html}</div>'
    
    elif node_type == 'dublinLogo':
        return '<div class="dublin-logo-container"><img src="data:image/svg+xml,<svg xmlns=\'http://www.w3.org/2000/svg\' viewBox=\'0 0 100 100\'><circle cx=\'50\' cy=\'50\' r=\'40\' fill=\'%232e8b57\'/><text x=\'50\' y=\'65\' text-anchor=\'middle\' fill=\'white\' font-size=\'36\'>☘</text></svg>" alt="Dublin City Logo" class="dublin-logo" width="80" height="80" /></div>'
    
    elif node_type == 'dublinTitle':
        level = attrs.get('level', 'main')
        text_content = ''.join(convert_tiptap_node_to_html(child) for child in content)
        return f'<h1 class="dublin-title dublin-title-{level}">{text_content}</h1>'
    
    elif node_type == 'sectionBreak':
        text = attrs.get('text', 'REGULAR MEETING 7:00 PM')
        return f'<div class="section-break">{text}</div>'
    
    elif node_type == 'noticeBox':
        title = attrs.get('title', 'Additional Meeting Procedures')
        content_html = ''.join(convert_tiptap_node_to_html(child) for child in content)
        return f'<div class="notice-box"><div class="notice-box-title">{title}</div><div class="notice-box-content">{content_html}</div></div>'
    
    elif node_type == 'text':
        return node.get('text', '')
    
    elif node_type == 'hardBreak':
        return '<br>'
    
    else:
        # For other node types, just convert their content
        return ''.join(convert_tiptap_node_to_html(child) for child in content)

def generate_sausalito_word_document(template_data):
    """Generate Sausalito Word-style document from HTML content"""
    try:
        # Get HTML content from TinyMCE
        html_content = template_data.get('html_content', '')
        title = template_data.get('title', 'Sausalito City Council Agenda')
        
        # Read the shared CSS file
        css_file_path = os.path.join(os.path.dirname(__file__), 'public', 'sausalito-agenda.css')
        sausalito_css = ""
        try:
            with open(css_file_path, 'r', encoding='utf-8') as css_file:
                sausalito_css = css_file.read()
        except FileNotFoundError:
            # Fallback CSS if file doesn't exist
            sausalito_css = """
                @page { size: letter; margin: 1in; }
                body { 
                    font-family: 'Times New Roman', Times, serif; 
                    font-size: 12pt; 
                    line-height: 1.4; 
                    color: #000000; 
                    background: white; 
                }
                h1, h2, h3 { font-weight: bold; margin-top: 20px; margin-bottom: 10px; }
                p { margin-bottom: 12px; }
                .section-heading { 
                    color: #1a365d; 
                    font-size: 16pt; 
                    text-transform: uppercase; 
                    border-bottom: 2px solid #1a365d; 
                    padding-bottom: 5px; 
                }
                .staff-report { 
                    background: #f7fafc; 
                    border-left: 4px solid #4299e1; 
                    padding: 15px; 
                    margin: 20px 0; 
                }
                .notice-box { 
                    border: 2px solid #e53e3e; 
                    background: #fed7d7; 
                    padding: 15px; 
                    margin: 20px 0; 
                    text-align: center; 
                }
                .page-break-before { page-break-before: always; }
            """
        
        # Generate complete HTML document
        full_html = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{title}</title>
            <style>
                {sausalito_css}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        return full_html
        
    except Exception as e:
        raise Exception(f"Sausalito Word document generation failed: {str(e)}")

def convert_html_to_pdf(html_content):
    """Convert HTML to PDF using WeasyPrint"""
    try:
        # Import weasyprint inside function to avoid pango errors
        import weasyprint
        
        # Create temporary file for PDF
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        
        # Generate PDF from HTML
        weasyprint.HTML(string=html_content).write_pdf(temp_file.name)
        
        return temp_file.name
    except Exception as e:
        raise Exception(f"HTML to PDF conversion failed: {str(e)}")

@app.route('/api/generate-pdf', methods=['POST'])
def generate_pdf():
    """Generate PDF from template data"""
    try:
        data = request.json
        template_type = data.get('template', 'sausalito-agenda')
        
        if template_type == 'dublin-agenda':
            # Generate HTML document for Dublin agenda
            html_content = generate_dublin_html_document(data)
            
            # Convert HTML to PDF
            pdf_file = convert_html_to_pdf(html_content)
            
            # Return PDF file
            return send_file(pdf_file, as_attachment=True, download_name='dublin_agenda.pdf')
        
        elif template_type == 'dublin-word':
            # Generate Dublin Word document
            docx_file = generate_dublin_word_document(data)
            
            # Convert to PDF
            pdf_file = convert_to_pdf(docx_file)
            
            # Clean up Word file
            os.unlink(docx_file)
            
            # Return PDF file
            return send_file(pdf_file, as_attachment=True, download_name='dublin_word_agenda.pdf')
        
        elif template_type == 'dublin-tiptap':
            # Generate Dublin TipTap document
            html_content = generate_dublin_tiptap_document(data)
            
            # Convert HTML to PDF
            pdf_file = convert_html_to_pdf(html_content)
            
            # Return PDF file
            return send_file(pdf_file, as_attachment=True, download_name='dublin_tiptap_agenda.pdf')
        
        elif template_type == 'sausalito-word':
            # Generate Sausalito Word-style document from HTML
            html_content = generate_sausalito_word_document(data)
            
            # Convert HTML to PDF
            pdf_file = convert_html_to_pdf(html_content)
            
            # Return PDF file
            return send_file(pdf_file, as_attachment=True, download_name='sausalito_word_agenda.pdf')
        
        else:
            # Generate Word document for agenda (Sausalito)
            docx_file = generate_word_document(data)
            
            # Convert to PDF
            pdf_file = convert_to_pdf(docx_file)
            
            # Clean up Word file
            os.unlink(docx_file)
            
            # Return PDF file
            return send_file(pdf_file, as_attachment=True, download_name='agenda.pdf')
        
    except Exception as e:
        return f"PDF generation failed: {str(e)}", 500

if __name__ == '__main__':
    # Get port from environment variable or default to 8000
    port = int(os.environ.get('FLASK_PORT', 8000))
    
    # Check if running in standalone mode (no Next.js frontend)
    standalone_mode = not os.path.exists('package.json')
    
    print("🚀 Starting Template Editor Demo Server")
    print("=" * 50)
    if standalone_mode:
        print(f"📍 Open your browser to: http://localhost:{port}")
        print("📝 Edit templates in real-time")
        print("👁️  See instant HTML preview")
        print("📄 Generate PDF on demand")
    else:
        print(f"🔧 Backend API running on port {port}")
        print("📄 PDF generation service ready")
        print("🔗 Frontend will connect automatically")
    print("=" * 50)
    
    app.run(debug=True, host='0.0.0.0', port=port) 